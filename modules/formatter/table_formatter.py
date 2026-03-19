"""表格格式模块 — 三线表、底纹、表内字号、跨页表头"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

    for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec is None:
            continue
        elem = tcBorders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(elem)
        elem.set(qn('w:val'), spec.get('val', 'single'))
        elem.set(qn('w:sz'), str(spec.get('sz', 4)))
        elem.set(qn('w:color'), spec.get('color', '000000'))
        elem.set(qn('w:space'), '0')


def apply_three_line_table(doc: Document, header_rows: int = 1) -> int:
    """将所有表格设为三线表格式"""
    count = 0
    for table in doc.tables:
        total_rows = len(table.rows)
        if total_rows == 0:
            continue

        thick = {'val': 'single', 'sz': '12', 'color': '000000'}  # 1.5磅
        thin = {'val': 'single', 'sz': '4', 'color': '000000'}    # 0.5磅
        none_border = {'val': 'none', 'sz': '0', 'color': '000000'}

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if i == 0:
                    _set_cell_border(cell, top=thick, bottom=thin, left=none_border, right=none_border)
                elif i == total_rows - 1:
                    _set_cell_border(cell, top=none_border, bottom=thick, left=none_border, right=none_border)
                elif i < header_rows:
                    _set_cell_border(cell, top=none_border, bottom=thin, left=none_border, right=none_border)
                else:
                    _set_cell_border(cell, top=none_border, bottom=none_border, left=none_border, right=none_border)

        count += 1
    return count


def set_table_font(doc: Document, cn_font: str = "宋体", en_font: str = "Times New Roman",
                   size_pt: float = 10.5) -> int:
    """统一表格内字体字号"""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(size_pt)
                        run.font.name = en_font
                        rpr = run._element.get_or_add_rPr()
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is None:
                            rfonts = run._element.makeelement(qn('w:rFonts'), {})
                            rpr.insert(0, rfonts)
                        rfonts.set(qn('w:eastAsia'), cn_font)
                        rfonts.set(qn('w:ascii'), en_font)
                        rfonts.set(qn('w:hAnsi'), en_font)
                        count += 1
    return count


def set_header_shading(doc: Document, color: str = "D9E2F3") -> int:
    """设置表头底纹颜色"""
    count = 0
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        for cell in table.rows[0].cells:
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
            count += 1
    return count


def clear_all_shading(doc: Document) -> int:
    """清除所有表格底纹"""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        tcPr.remove(shd)
                        count += 1
    return count


def enable_repeat_header_row(doc: Document) -> int:
    """跨页表格自动重复表头行"""
    count = 0
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        row = table.rows[0]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        existing = trPr.find(qn('w:tblHeader'))
        if existing is None:
            tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
            trPr.append(tbl_header)
            count += 1
    return count
