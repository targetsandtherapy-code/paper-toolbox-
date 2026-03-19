"""页面设置模块 — 页边距、页眉页脚、页码"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from modules.formatter.template_parser import TemplateConfig


def set_page_margins(doc: Document, config: TemplateConfig) -> int:
    """按模板设置页边距"""
    if not config.pages:
        return 0
    page = config.pages[0]
    count = 0
    for sec in doc.sections:
        sec.top_margin = Cm(page.top_cm)
        sec.bottom_margin = Cm(page.bottom_cm)
        sec.left_margin = Cm(page.left_cm)
        sec.right_margin = Cm(page.right_cm)
        if page.header_distance_cm:
            sec.header_distance = Cm(page.header_distance_cm)
        if page.footer_distance_cm:
            sec.footer_distance = Cm(page.footer_distance_cm)
        count += 1
    return count


def set_page_margins_custom(doc: Document, top: float = 2.5, bottom: float = 2.5,
                            left: float = 3.0, right: float = 2.0) -> int:
    """自定义页边距（cm）"""
    count = 0
    for sec in doc.sections:
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)
        count += 1
    return count


def set_header_text(doc: Document, text: str, font_name: str = "宋体",
                    font_size_pt: float = 10.5, section_index: int = -1):
    """设置页眉文字"""
    sections = doc.sections if section_index < 0 else [doc.sections[section_index]]
    for sec in sections:
        header = sec.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            p = header.paragraphs[0]
        else:
            p = header.add_paragraph()

        p.clear()
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = run._element.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), font_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_header_bottom_border(p)


def _add_header_bottom_border(paragraph):
    """给页眉段落添加下边框线"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = paragraph._element.makeelement(qn('w:pBdr'), {})
        pPr.append(pBdr)

    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = paragraph._element.makeelement(qn('w:bottom'), {})
        pBdr.append(bottom)

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')


def remove_first_section_header(doc: Document):
    """移除第一个 Section 的页眉（封面页不显示页眉）"""
    if doc.sections:
        sec = doc.sections[0]
        sec.different_first_page_header_footer = True
        first_header = sec.first_page_header
        for p in first_header.paragraphs:
            p.clear()
