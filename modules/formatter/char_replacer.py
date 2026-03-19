"""字符替换模块 — 标点互换、全角半角、空格空行清理、中英文间距"""
import re
from docx import Document

# 中文标点 → 英文标点
CN_TO_EN = {
    "\uff0c": ",", "\u3002": ".", "\uff1b": ";", "\uff1a": ":",
    "\uff01": "!", "\uff1f": "?", "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"', "\uff08": "(", "\uff09": ")",
    "\u3010": "[", "\u3011": "]", "\u3001": ",", "\u2014\u2014": "--",
}

EN_TO_CN = {v: k for k, v in CN_TO_EN.items() if k != "\u2014\u2014"}
EN_TO_CN["--"] = "\u2014\u2014"


def replace_punctuation(doc: Document, direction: str = "cn_to_en") -> int:
    """替换标点符号
    direction: 'cn_to_en' 中文→英文 | 'en_to_cn' 英文→中文
    """
    mapping = CN_TO_EN if direction == "cn_to_en" else EN_TO_CN
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            new_text = original
            for old, new in mapping.items():
                new_text = new_text.replace(old, new)
            if new_text != original:
                run.text = new_text
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original = run.text
                        new_text = original
                        for old, new in mapping.items():
                            new_text = new_text.replace(old, new)
                        if new_text != original:
                            run.text = new_text
                            count += 1
    return count


def clean_extra_spaces(doc: Document) -> int:
    """清理段首空格和多余连续空格"""
    count = 0
    for para in doc.paragraphs:
        if not para.runs:
            continue
        first_run = para.runs[0]
        stripped = first_run.text.lstrip()
        if stripped != first_run.text:
            first_run.text = stripped
            count += 1
        for run in para.runs:
            cleaned = re.sub(r"[ \t]{2,}", " ", run.text)
            if cleaned != run.text:
                run.text = cleaned
                count += 1
    return count


def clean_empty_paragraphs(doc: Document, max_consecutive: int = 1) -> int:
    """清理连续空段落，保留最多 max_consecutive 个"""
    removed = 0
    consecutive = 0
    to_remove = []

    for para in doc.paragraphs:
        if not para.text.strip():
            consecutive += 1
            if consecutive > max_consecutive:
                to_remove.append(para)
        else:
            consecutive = 0

    for para in to_remove:
        p_element = para._element
        p_element.getparent().remove(p_element)
        removed += 1

    return removed


def add_space_between_cn_en(doc: Document) -> int:
    """中英文/中文数字之间自动加空格"""
    pattern_cn_en = re.compile(r'([\u4e00-\u9fff])([A-Za-z0-9])')
    pattern_en_cn = re.compile(r'([A-Za-z0-9])([\u4e00-\u9fff])')

    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            new_text = pattern_cn_en.sub(r'\1 \2', original)
            new_text = pattern_en_cn.sub(r'\1 \2', new_text)
            if new_text != original:
                run.text = new_text
                count += 1
    return count


def add_space_between_number_unit(doc: Document) -> int:
    """数字与单位之间加空格（如 100mL → 100 mL），但℃等特殊单位除外"""
    pattern = re.compile(r'(\d)([a-zA-Z\u03bc])')  # μ 也包含
    exclude_after = re.compile(r'(\d)([\u2103])') # ℃ 不加空格

    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            new_text = pattern.sub(r'\1 \2', original)
            if new_text != original:
                run.text = new_text
                count += 1
    return count
