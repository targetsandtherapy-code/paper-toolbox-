"""段落格式模块 -- 行距、缩进、段前段后、对齐"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from modules.formatter.template_parser import TemplateConfig
from modules.formatter.font_formatter import _get_heading_level, _detect_heading_by_pattern


_SECTION_TITLES = frozenset({
    "内容摘要", "摘要", "摘 要", "abstract",
    "目录", "目  录", "目 录",
    "参考文献", "参 考 文 献", "references",
    "致谢", "致 谢", "致  谢", "acknowledgments",
    "附录", "附 录",
})


def _is_section_title(text: str) -> bool:
    """判断是否为独立成段的章节标题（应居中）"""
    return text.strip().lower() in _SECTION_TITLES


def format_paragraphs(doc: Document, config: TemplateConfig) -> dict:
    """按模板配置格式化段落"""
    stats = {"body": 0, "heading": 0, "section_title": 0}

    heading_specs = {
        1: config.heading1,
        2: config.heading2,
        3: config.heading3,
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        level = _get_heading_level(para) or _detect_heading_by_pattern(text)

        if _is_section_title(text):
            spec = config.abstract_title
            if spec:
                if spec.line_spacing:
                    pf.line_spacing = spec.line_spacing
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                if spec.space_before_pt:
                    pf.space_before = Pt(spec.space_before_pt)
                if spec.space_after_pt:
                    pf.space_after = Pt(spec.space_after_pt)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            stats["section_title"] += 1

        elif level and level in heading_specs and heading_specs[level]:
            spec = heading_specs[level]
            if spec.line_spacing:
                pf.line_spacing = spec.line_spacing
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            elif spec.line_spacing_pt:
                pf.line_spacing = Pt(spec.line_spacing_pt)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            if spec.space_before_pt:
                pf.space_before = Pt(spec.space_before_pt)
            if spec.space_after_pt:
                pf.space_after = Pt(spec.space_after_pt)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            stats["heading"] += 1

        elif config.body_text and not level:
            spec = config.body_text
            if spec.line_spacing:
                pf.line_spacing = spec.line_spacing
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            elif spec.line_spacing_pt:
                pf.line_spacing = Pt(spec.line_spacing_pt)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            if spec.first_indent_pt:
                pf.first_line_indent = Pt(spec.first_indent_pt)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            stats["body"] += 1

    return stats


def set_body_line_spacing(doc: Document, spacing_pt: float = 20.0) -> int:
    """统一设置正文行距（固定值）"""
    count = 0
    for para in doc.paragraphs:
        if para.text.strip() and len(para.text.strip()) > 5:
            level = _get_heading_level(para) or _detect_heading_by_pattern(para.text.strip())
            if not level and not _is_section_title(para.text.strip()):
                para.paragraph_format.line_spacing = Pt(spacing_pt)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                count += 1
    return count


def set_first_line_indent(doc: Document, indent_chars: float = 2.0, font_size_pt: float = 12.0) -> int:
    """设置正文首行缩进"""
    indent_pt = indent_chars * font_size_pt
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 5:
            level = _get_heading_level(para) or _detect_heading_by_pattern(text)
            if not level and not _is_section_title(text):
                para.paragraph_format.first_line_indent = Pt(indent_pt)
                count += 1
    return count
