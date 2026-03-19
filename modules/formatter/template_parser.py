"""模板解析器 — 从学校官方 .docx 模板中自动提取格式参数"""
from dataclasses import dataclass, field
from typing import Optional
from docx import Document
from docx.shared import Pt, Cm, Emu


@dataclass
class FontSpec:
    cn_name: Optional[str] = None  # 中文字体名
    en_name: Optional[str] = None  # 英文字体名
    size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None

    def __repr__(self):
        parts = []
        if self.cn_name:
            parts.append(self.cn_name)
        if self.en_name and self.en_name != self.cn_name:
            parts.append(self.en_name)
        if self.size_pt:
            parts.append(f"{self.size_pt}pt")
        if self.bold:
            parts.append("Bold")
        return "/".join(parts) if parts else "inherit"


@dataclass
class ParagraphSpec:
    font: FontSpec = field(default_factory=FontSpec)
    line_spacing: Optional[float] = None      # 倍数行距 (如 1.5)
    line_spacing_pt: Optional[float] = None   # 磅值行距 (如 20)
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    first_indent_pt: Optional[float] = None
    alignment: Optional[str] = None           # LEFT/CENTER/RIGHT/JUSTIFY


@dataclass
class PageSpec:
    width_cm: float = 21.0
    height_cm: float = 29.7
    top_cm: float = 2.5
    bottom_cm: float = 2.5
    left_cm: float = 3.0
    right_cm: float = 2.0
    header_distance_cm: float = 2.0
    footer_distance_cm: float = 2.0
    header_text: str = ""


@dataclass
class TemplateConfig:
    """从模板中提取的完整格式配置"""
    pages: list[PageSpec] = field(default_factory=list)
    heading1: Optional[ParagraphSpec] = None
    heading2: Optional[ParagraphSpec] = None
    heading3: Optional[ParagraphSpec] = None
    body_text: Optional[ParagraphSpec] = None
    abstract_title: Optional[ParagraphSpec] = None
    abstract_body: Optional[ParagraphSpec] = None
    keywords: Optional[ParagraphSpec] = None
    toc_title: Optional[ParagraphSpec] = None
    cover_fields: list[ParagraphSpec] = field(default_factory=list)
    named_styles: dict[str, ParagraphSpec] = field(default_factory=dict)

    def summary(self) -> str:
        lines = ["=== 模板格式配置 ==="]
        if self.pages:
            p = self.pages[0]
            lines.append(f"页面: 上{p.top_cm}cm 下{p.bottom_cm}cm 左{p.left_cm}cm 右{p.right_cm}cm")
            lines.append(f"页眉距{p.header_distance_cm}cm 页脚距{p.footer_distance_cm}cm")
        if self.heading1:
            lines.append(f"一级标题: {self.heading1.font}")
        if self.heading2:
            lines.append(f"二级标题: {self.heading2.font}")
        if self.heading3:
            lines.append(f"三级标题: {self.heading3.font}")
        if self.body_text:
            lines.append(f"正文: {self.body_text.font} 行距={self.body_text.line_spacing or self.body_text.line_spacing_pt}")
        if self.abstract_title:
            lines.append(f"摘要标题: {self.abstract_title.font}")
        if self.keywords:
            lines.append(f"关键词: {self.keywords.font}")
        return "\n".join(lines)


def _emu_to_cm(emu) -> float:
    if emu is None:
        return 0.0
    return round(emu / 360000, 2)


def _emu_to_pt(emu) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / 12700, 1)


def _extract_font(run) -> FontSpec:
    f = run.font
    cn_name = None
    en_name = f.name

    from docx.oxml.ns import qn
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            cn_name = rfonts.get(qn('w:eastAsia'))
            if not en_name:
                en_name = rfonts.get(qn('w:ascii')) or rfonts.get(qn('w:hAnsi'))

    return FontSpec(
        cn_name=cn_name or en_name,
        en_name=en_name,
        size_pt=_emu_to_pt(f.size) if f.size else None,
        bold=f.bold,
        italic=f.italic,
    )


def _extract_paragraph_spec(para) -> ParagraphSpec:
    pf = para.paragraph_format

    font = FontSpec()
    if para.runs:
        font = _extract_font(para.runs[0])

    line_spacing = None
    line_spacing_pt = None
    if pf.line_spacing is not None:
        if pf.line_spacing > 50000:
            line_spacing_pt = _emu_to_pt(pf.line_spacing)
        else:
            line_spacing = round(pf.line_spacing, 2)

    alignment = None
    if pf.alignment is not None:
        alignment = str(pf.alignment).split(".")[-1].split("(")[0].strip()

    return ParagraphSpec(
        font=font,
        line_spacing=line_spacing,
        line_spacing_pt=line_spacing_pt,
        space_before_pt=_emu_to_pt(pf.space_before),
        space_after_pt=_emu_to_pt(pf.space_after),
        first_indent_pt=_emu_to_pt(pf.first_line_indent),
        alignment=alignment,
    )


def _is_abstract_title(text: str) -> bool:
    t = text.strip().lower()
    return t in ("内容摘要", "摘要", "摘 要", "abstract", "中文摘要", "英文摘要")


def _is_keywords(text: str) -> bool:
    t = text.strip().lower()
    return t.startswith("关键词") or t.startswith("key words") or t.startswith("keywords")


def _is_toc_title(text: str) -> bool:
    t = text.strip()
    return t in ("目录", "目  录", "目 录", "TABLE OF CONTENTS", "CONTENTS")


def parse_template(docx_path: str) -> TemplateConfig:
    """解析 Word 模板文件，提取格式配置"""
    doc = Document(docx_path)
    config = TemplateConfig()

    for sec in doc.sections:
        page = PageSpec(
            width_cm=_emu_to_cm(sec.page_width),
            height_cm=_emu_to_cm(sec.page_height),
            top_cm=_emu_to_cm(sec.top_margin),
            bottom_cm=_emu_to_cm(sec.bottom_margin),
            left_cm=_emu_to_cm(sec.left_margin),
            right_cm=_emu_to_cm(sec.right_margin),
            header_distance_cm=_emu_to_cm(sec.header_distance),
            footer_distance_cm=_emu_to_cm(sec.footer_distance),
        )
        if sec.header:
            for p in sec.header.paragraphs:
                if p.text.strip():
                    page.header_text = p.text.strip()
        config.pages.append(page)

    # 提取样式定义
    for style in doc.styles:
        if not hasattr(style, "font") or style.font is None:
            continue
        name = style.name
        f = style.font
        if name.startswith("Heading 1") or name == "标题 1":
            config.heading1 = ParagraphSpec(font=FontSpec(
                cn_name=f.name, en_name=f.name,
                size_pt=_emu_to_pt(f.size) if f.size else None,
                bold=f.bold,
            ))
        elif name.startswith("Heading 2") or name == "标题 2":
            config.heading2 = ParagraphSpec(font=FontSpec(
                cn_name=f.name, en_name=f.name,
                size_pt=_emu_to_pt(f.size) if f.size else None,
                bold=f.bold,
            ))
        elif name.startswith("Heading 3") or name == "标题 3":
            config.heading3 = ParagraphSpec(font=FontSpec(
                cn_name=f.name, en_name=f.name,
                size_pt=_emu_to_pt(f.size) if f.size else None,
                bold=f.bold,
            ))

    # 从段落内容推断更多格式
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if _is_abstract_title(text) and config.abstract_title is None:
            config.abstract_title = _extract_paragraph_spec(para)

        elif _is_keywords(text) and config.keywords is None:
            config.keywords = _extract_paragraph_spec(para)

        elif _is_toc_title(text) and config.toc_title is None:
            config.toc_title = _extract_paragraph_spec(para)

        # 尝试找正文段落（较长的、非标题的段落）
        elif len(text) > 80 and config.body_text is None:
            spec = _extract_paragraph_spec(para)
            if spec.font.size_pt and 11 <= spec.font.size_pt <= 13:
                config.body_text = spec

    # 如果没从段落中找到标题样式的格式细节，用样式定义补充
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name.startswith("Heading 1") and config.heading1:
            full = _extract_paragraph_spec(para)
            if full.space_before_pt:
                config.heading1.space_before_pt = full.space_before_pt
            if full.space_after_pt:
                config.heading1.space_after_pt = full.space_after_pt
            if full.line_spacing:
                config.heading1.line_spacing = full.line_spacing
            break

    return config
