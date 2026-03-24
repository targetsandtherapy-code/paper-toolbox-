"""Word 文档解析 & 角标识别模块"""
import re
from dataclasses import dataclass, field
from pathlib import Path
from docx import Document


@dataclass
class CitationMarker:
    """单个引用角标"""
    ids: list[int]           # 角标编号列表，如 [1] → [1], [1,2] → [1,2], [1-3] → [1,2,3]
    paragraph_text: str      # 角标所在段落全文
    context_before: str      # 角标所在完整句（上一句末标点至本句末标点，含角标）
    paragraph_index: int     # 段落在文档中的位置索引
    raw_marker: str          # 原始角标文本，如 "[1]", "[1-3]"


MARKER_PATTERN = re.compile(
    r'\[(\d+(?:\s*[-–—]\s*\d+)?(?:\s*[,，]\s*\d+(?:\s*[-–—]\s*\d+)?)*)\]'
)


def _expand_marker_ids(raw: str) -> list[int]:
    """展开角标编号: "[1-3,5]" → [1,2,3,5]"""
    raw = raw.strip("[]")
    ids = []
    for part in re.split(r'[,，]', raw):
        part = part.strip()
        range_match = re.match(r'(\d+)\s*[-–—]\s*(\d+)', part)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            ids.extend(range(start, end + 1))
        elif part.isdigit():
            ids.append(int(part))
    return sorted(set(ids))


# 句末：中文仅「。」、英文仅「.」；分号/叹号/问号等不作为句界（避免半角 ; 误切）
_SENTENCE_END_RE = re.compile(r"[。.]")


def _extract_sentence_containing_marker(
    text: str,
    marker_start: int,
    marker_end: int,
    max_chars: int = 2000,
) -> str:
    """提取角标所在整句：从**上一句**的句末（仅「。」或「.」）之后起，到**本句**的「。」或「.」为止（含角标及角标后同句文字）。

    分号、逗号、叹号、问号、换行均不作为句界。若本句无后续「。」/「.」（如段末），则取到段落末尾。过长时以角标为中心截断到 max_chars。
    """
    before = text[:marker_start]
    break_ends = [m.end() for m in _SENTENCE_END_RE.finditer(before)]
    sent_start = break_ends[-1] if break_ends else 0

    after = text[marker_end:]
    m = _SENTENCE_END_RE.search(after)
    sent_end = marker_end + m.end() if m else len(text)

    sent = text[sent_start:sent_end].strip()
    if not sent:
        return ""

    if len(sent) > max_chars:
        rel = marker_start - sent_start
        half = max_chars // 2
        lo = max(0, min(rel - half, len(sent) - max_chars))
        sent = sent[lo : lo + max_chars].strip()

    return sent


class DocParser:
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if self.file_path.suffix.lower() != ".docx":
            raise ValueError(f"仅支持 .docx 格式，当前: {self.file_path.suffix}")
        self.doc = Document(str(self.file_path))

    _REF_SECTION_HEADERS = re.compile(
        r"^(参\s*考\s*文\s*献|references?|bibliography)\s*$", re.I
    )

    def _iter_body_paragraphs(self):
        """遍历主文档 body 子树内全部 ``w:p``，返回 (段落序号, 段落纯文本)。

        使用 ``body.iter(w:p)`` 深度遍历，覆盖：
        - 表格单元格内段落（含合并单元格等 python-docx ``Table.rows`` 无法完整展开的情况）；
        - 内容控件 ``w:sdt``、嵌套块等下的段落。

        不再按「body 直接子节点 + Table API」拼接，避免漏段导致角标数量偏少。
        """
        from docx.oxml.ns import qn

        body = self.doc.element.body
        idx = 0
        for para in body.iter(qn("w:p")):
            text = "".join(node.text or "" for node in para.iter(qn("w:t")))
            yield idx, text
            idx += 1

    def extract_markers(self) -> list[CitationMarker]:
        """提取文档正文中的引用角标（在最后一个「参考文献」标题处停止）。
        同时扫描表格单元格中的段落。"""
        all_paras: list[tuple[int, str]] = list(self._iter_body_paragraphs())

        last_ref_idx = -1
        for i, (vidx, text) in enumerate(all_paras):
            if self._REF_SECTION_HEADERS.match(text.strip()):
                last_ref_idx = i

        markers = []
        for seq, (vidx, text) in enumerate(all_paras):
            text = text.strip()
            if not text:
                continue
            if seq == last_ref_idx:
                break

            for match in MARKER_PATTERN.finditer(text):
                raw_marker = match.group(0)
                ids = _expand_marker_ids(raw_marker)
                context_before = _extract_sentence_containing_marker(
                    text, match.start(), match.end()
                )

                markers.append(CitationMarker(
                    ids=ids,
                    paragraph_text=text,
                    context_before=context_before,
                    paragraph_index=vidx,
                    raw_marker=raw_marker,
                ))

        return markers

    def extract_markers_grouped(self) -> dict[int, CitationMarker]:
        """按角标编号分组返回（每个编号取第一次出现的位置）"""
        markers = self.extract_markers()
        grouped = {}
        for m in markers:
            for cid in m.ids:
                if cid not in grouped:
                    grouped[cid] = m
        return dict(sorted(grouped.items()))

    _SKIP_TITLES = {
        "摘要", "摘  要", "abstract", "目录", "致谢", "参考文献", "附录",
        "关键词", "keywords", "目  录",
    }
    _SKIP_TITLE_PATTERNS = re.compile(
        r"^[（(]?\s*(硕\s*士\s*学\s*位\s*论\s*文|博\s*士\s*学\s*位\s*论\s*文|"
        r"学\s*位\s*论\s*文|毕\s*业\s*论\s*文|专\s*业\s*学\s*位|学\s*术\s*学\s*位|"
        r"master['\u2019]?s?\s*thesis|doctoral\s*dissertation|"
        r"学\s*号|姓\s*名|导\s*师|指\s*导\s*教\s*师|专\s*业|学\s*院|"
        r"研\s*究\s*生|学\s*位\s*类\s*别|学\s*科\s*门\s*类|"
        r"分类号|密级|UDC|申请人|答辩日期|培养单位)\s*[)）]?\s*[:：]?\s*$",
        re.I,
    )

    def _is_skip_title(self, text: str) -> bool:
        normalized = text.lower().replace(" ", "").replace("\u3000", "")
        if normalized in {s.replace(" ", "") for s in self._SKIP_TITLES}:
            return True
        if self._SKIP_TITLE_PATTERNS.match(text.strip()):
            return True
        stripped = re.sub(r"[（()）\s]", "", text)
        if stripped in ("专业学位", "学术学位", "全日制", "非全日制", "在职"):
            return True
        return False

    def get_title(self) -> str:
        """提取论文标题（取前几个非空段落中最可能是标题的那个）"""
        for para in self.doc.paragraphs[:20]:
            text = para.text.strip()
            if not text or len(text) < 4:
                continue
            if self._is_skip_title(text):
                continue
            if MARKER_PATTERN.search(text):
                continue
            if para.style and para.style.name and 'title' in para.style.name.lower():
                return text
        for para in self.doc.paragraphs[:20]:
            text = para.text.strip()
            if not text or len(text) < 6 or len(text) > 80:
                continue
            if self._is_skip_title(text):
                continue
            if MARKER_PATTERN.search(text):
                continue
            if text.startswith("关键词") or text.startswith("Keywords"):
                continue
            return text
        return ""

    def get_full_text(self) -> str:
        """获取文档全文"""
        return "\n".join(p.text for p in self.doc.paragraphs if p.text.strip())

    def get_paragraphs(self) -> list[str]:
        """获取所有非空段落"""
        return [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
