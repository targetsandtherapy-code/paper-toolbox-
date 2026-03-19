"""Streamlit 页面 - 查重预检"""
import sys
import os
import re
import tempfile
from pathlib import Path

os.environ["PYTHONIOENCODING"] = "utf-8"
sys.path.insert(0, str(Path(__file__).parent.parent))

import streamlit as st

st.title("🔍 查重预检")
st.markdown("在提交正式查重前，先进行自查，标记高风险段落。")

mode = st.radio("检测模式", ["文档内部自查", "两篇文档比对"], horizontal=True)

if mode == "文档内部自查":
    st.markdown("检测论文内部是否存在大段重复（自我复制/结构雷同）")

    upload_tab, text_tab = st.tabs(["上传文档", "粘贴文本"])
    with upload_tab:
        file_a = st.file_uploader("上传论文 (.docx)", type=["docx"], key="self_check")
    with text_tab:
        text_a = st.text_area("粘贴论文文本", height=300, key="self_text")

    threshold = st.slider("相似度阈值", 0.5, 1.0, 0.85, 0.05,
        help="越高越严格，0.85 表示 85% 以上相似才标记")

    if st.button("🔍 开始自查", type="primary", use_container_width=True):
        full_text = ""
        if file_a:
            from docx import Document
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp.write(file_a.getvalue())
            tmp.close()
            doc = Document(tmp.name)
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif text_a and text_a.strip():
            full_text = text_a.strip()
        else:
            st.error("请上传文档或粘贴文本")
            st.stop()

        with st.spinner("正在分析..."):
            sentences = re.split(r'(?<=[。！？.!?])\s*', full_text)
            sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) >= 10]

            from modules.checker.simhash import self_check
            from modules.checker.text_compare import find_repeated_sentences

            duplicates_sh = self_check(sentences, threshold=threshold)
            duplicates_ng = find_repeated_sentences(full_text, threshold=threshold)

        all_dups = []
        seen = set()
        for d in duplicates_sh:
            key = (d["sentence_a"][:30], d["sentence_b"][:30])
            if key not in seen:
                seen.add(key)
                all_dups.append(d)
        for d in duplicates_ng:
            key = (d["sentence_a"][:30], d["sentence_b"][:30])
            if key not in seen:
                seen.add(key)
                all_dups.append({"sentence_a": d["sentence_a"], "sentence_b": d["sentence_b"],
                                 "similarity": d["overlap"]})

        st.subheader("检测结果")
        st.metric("总句子数", len(sentences))

        if all_dups:
            st.warning(f"发现 {len(all_dups)} 组相似段落")
            for i, d in enumerate(all_dups):
                with st.container(border=True):
                    st.markdown(f"**相似度: {d['similarity']:.0%}**")
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown("**段落 A:**")
                        st.text(d.get("sentence_a", "")[:200])
                    with c2:
                        st.markdown("**段落 B:**")
                        st.text(d.get("sentence_b", "")[:200])
        else:
            st.success("未发现内部重复段落")

elif mode == "两篇文档比对":
    st.markdown("比对两篇文档的文本相似度")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**文档 A（你的论文）**")
        file_a = st.file_uploader("上传文档 A (.docx)", type=["docx"], key="cmp_a")
        text_a = st.text_area("或粘贴文本 A", height=200, key="cmp_text_a")
    with c2:
        st.markdown("**文档 B（参考文档）**")
        file_b = st.file_uploader("上传文档 B (.docx)", type=["docx"], key="cmp_b")
        text_b = st.text_area("或粘贴文本 B", height=200, key="cmp_text_b")

    if st.button("🔍 开始比对", type="primary", use_container_width=True):
        def get_text(file_obj, text_input):
            if file_obj:
                from docx import Document
                tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                tmp.write(file_obj.getvalue())
                tmp.close()
                doc = Document(tmp.name)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif text_input and text_input.strip():
                return text_input.strip()
            return ""

        ta = get_text(file_a, text_a)
        tb = get_text(file_b, text_b)

        if not ta or not tb:
            st.error("请提供两篇文档的文本")
            st.stop()

        with st.spinner("正在比对..."):
            from modules.checker.text_compare import compute_document_similarity, highlight_repeated_segments

            result = compute_document_similarity(ta, tb)
            segments = highlight_repeated_segments(ta, tb)

        st.subheader("比对结果")

        ratio = result["overlap_ratio"]
        color = "🟢" if ratio < 0.15 else "🟡" if ratio < 0.30 else "🔴"
        st.metric(f"{color} 整体重复率", f"{ratio:.1%}")
        st.caption(f"重复片段: {result['overlap_count']} / 总片段: {result['total_a']}")

        if segments:
            st.markdown("### 重复片段详情")
            for i, seg in enumerate(segments[:30]):
                with st.container(border=True):
                    st.markdown(f"**片段 {i+1}** ({len(seg['text'])} 字)")
                    st.text(seg["text"][:200])
        elif ratio > 0:
            st.info("重复内容分散在多个短片段中")
        else:
            st.success("未检测到重复内容")
